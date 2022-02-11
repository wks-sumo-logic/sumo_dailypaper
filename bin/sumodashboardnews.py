#!/usr/bin/env python3

import os
import sys
import fnmatch
import pathlib
import configparser
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

NOWTIME = datetime.now().strftime("%Y/%m/%d, %H:%M:%S")

DATE = datetime.now().strftime("%Y%m%d")
TIME = datetime.now().strftime("%H%M%S")

REPORTNAME = 'simple.' + DATE + '.docx'

JPGLIST = list()

CFGFILE = sys.argv[1]
print(CFGFILE)

CONFIG = configparser.ConfigParser()
CONFIG.optionxform = str
CONFIG.read(CFGFILE)

DEFAULTDICT = dict(CONFIG.items('Default'))
DASHBOARDDICT = dict(CONFIG.items('Dashboards'))

EXPORTDIR = CONFIG.get("Default", "EXPORTDIR")
SRCDIR = os.path.abspath(EXPORTDIR)

OUTPUTDIR = CONFIG.get("Default", "OUTPUTDIR")
OUTPUTFILE = '.'.join((CONFIG.get("Default", "OUTPUTFILE"), DATE, 'docx' ) ) 

NEWSPAPER = os.path.abspath(os.path.join(OUTPUTDIR, OUTPUTFILE))

document = Document()

header_section = document.sections[0]
header = header_section.header
header_text = header.paragraphs[0]
header_text.text = 'Report: {} Generated: {}'.format(REPORTNAME, NOWTIME)

footer_section = document.sections[0]
footer = footer_section.footer
footer_text = footer.paragraphs[0]
header_text.text = 'Courtesy of SumoDashboardNews!'

style = document.styles['Normal']
font = style.font

font.name = 'Calibri'
font.size = Pt(14)


COUNTER = 1
for key, value in DASHBOARDDICT.items():
    matchname = fnmatch.filter(os.listdir(SRCDIR), key + '.*.jpg')[0]
    if matchname is not None:
        jpgname = os.path.abspath(os.path.join(SRCDIR, matchname))
        picture_title = 'Title {}: {}'.format(COUNTER, value)
        document.add_heading(picture_title, 0)
        document.add_picture(jpgname, width=Inches(4))
        paragraph = document.add_paragraph('This is a simple summary; this could include a table with RAG values')
        document.add_page_break()
        COUNTER = COUNTER + 1

document.save(NEWSPAPER)
