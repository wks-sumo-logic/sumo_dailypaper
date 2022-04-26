#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Exaplanation: sumologic_dashboard_news makes a newspaper from dashboard exports

Usage:
    $ python  sumologic_dashboard_news [ options ]

Style:
    Google Python Style Guide:
    http://google.github.io/styleguide/pyguide.html

    @name           sumologic_dashboard_news
    @version        2.50
    @author-name    Wayne Schmidt
    @author-email   wschmidt@sumologic.com
    @license-name   Apache
    @license-url    https://www.apache.org/licenses/LICENSE-2.0
"""

__version__ = 2.50
__author__ = "Wayne Schmidt (wschmidt@sumologic.com)"

import json
import os
import sys
import argparse
import time
import fnmatch
import configparser
from datetime import datetime
import tzlocal
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import pdf2image
import requests

try:
    import cookielib
except ImportError:
    import http.cookiejar as cookielib

sys.dont_write_bytecode = 1

MY_CFG = 'undefined'
PARSER = argparse.ArgumentParser(description="""

sumologic_dashboard_news is a tool to build your own newpaper from dashboard output

""")

PARSER.add_argument("-a", metavar='<secret>', dest='MY_SECRET', \
                    help="set query authkey (format: <key>:<secret>) ")

PARSER.add_argument("-d", metavar='<dashboard>', dest='DASHBOARDLIST', \
                    action='append', help="set dashboard uid (list format)")

PARSER.add_argument("-c", metavar='<configfile>', dest='CONFIG', \
                    required=True,help="set config file")

PARSER.add_argument("-s", metavar='<sleeptime>', default=2, dest='SLEEPTIME', \
                    help="set sleep time to check results")

PARSER.add_argument("-v", type=int, default=0, metavar='<verbose>', \
                    dest='verbose', help="increase verbosity")

ARGS = PARSER.parse_args()

SLEEP = ARGS.SLEEPTIME

NOWTIME = datetime.now().strftime("%Y/%m/%d, %H:%M:%S")

DATE = datetime.now().strftime("%Y%m%d")

TIME = datetime.now().strftime("%H%M%S")

CONFIG = configparser.ConfigParser()
CONFIG.optionxform = str
CONFIG.read(ARGS.CONFIG)

def resolve_option_variables():
    """
    Validates and confirms all necessary variables for the script
    """

    if ARGS.MY_SECRET:
        (keyname, keysecret) = ARGS.MY_SECRET.split(':')
        os.environ['SUMO_UID'] = keyname
        os.environ['SUMO_KEY'] = keysecret

def resolve_config_variables():
    """
    Validates and confirms all necessary variables for the script
    """

    if ARGS.CONFIG:
        cfgfile = os.path.abspath(ARGS.CONFIG)
        configobj = configparser.ConfigParser()
        configobj.optionxform = str
        configobj.read(cfgfile)

        if ARGS.verbose > 8:
            print('Displaying Config Contents:')
            print(dict(configobj.items('Default')))

        if configobj.has_option("Default", "SUMO_UID"):
            os.environ['SUMO_UID'] = configobj.get("Default", "SUMO_UID")

        if configobj.has_option("Default", "SUMO_KEY"):
            os.environ['SUMO_KEY'] = configobj.get("Default", "SUMO_KEY")

def initialize_variables():
    """
    Validates and confirms all necessary variables for the script
    """

    resolve_option_variables()

    resolve_config_variables()

    try:
        my_uid = os.environ['SUMO_UID']
        my_key = os.environ['SUMO_KEY']

    except KeyError as myerror:
        print('Environment Variable Not Set :: {} '.format(myerror.args[0]))

    return my_uid, my_key

( sumo_uid, sumo_key ) = initialize_variables()

DEFAULTDICT = dict(CONFIG.items('Default'))
DASHBOARDDICT = dict(CONFIG.items('Dashboards'))

EXPORTDIR = '/var/tmp/dashboardexport'

REPORTDIR = '/var/tmp/dashboardnews'

REPORTTAG = 'sumodashboardnews'

def main():
    """
    Wrapper for the newspaper export, creation, and optional publishing
    """

    prepare_environment(EXPORTDIR,REPORTDIR)

    export_dashboards(sumo_uid, sumo_key, EXPORTDIR, DASHBOARDDICT)

    create_newspaper(EXPORTDIR,REPORTDIR)

def prepare_environment(exportdir,reportdir):
    """
    Build the environment
    """

    if ARGS.verbose > 5:
        print('Creating: {}'.format(exportdir))
    os.makedirs(exportdir, exist_ok=True)

    if ARGS.verbose > 5:
        print('Creating: {}'.format(reportdir))
    os.makedirs(reportdir, exist_ok=True)

def export_dashboards(sumouid, sumokey, exportdir, dashboarddict):
    """
    Export the dashboards based on a dictionary
    """

    exporter=SumoApiClient(sumouid, sumokey)

    tzname = tzlocal.get_localzone().zone

    for dashboardkey, _dashboardname in dashboarddict.items():
        export = exporter.run_export_job(dashboardkey,timezone=tzname,exportFormat='Pdf')

        if export['status'] != 'Success':
            print('Job: {} Status: {}'.format({export['job']}, {export['status']}))
        else:
            outputfile = "{dir}/{file}.{ext}".format(dir=exportdir,file=dashboardkey,ext='pdf')

            if ARGS.verbose > 3:
                print('Writing File: {}'.format(outputfile))

            with open(outputfile, "wb") as fileobject:
                fileobject.write(export['bytes'])

    for path in os.listdir(exportdir):
        file_name = os.path.join(exportdir, path)
        if ARGS.verbose > 3:
            print('Converting File: {}'.format(file_name))
        if os.path.isfile(file_name):
            extension = os.path.splitext(file_name)[1]
            if extension == '.pdf':
                images = pdf2image.convert_from_path(file_name)
                for i in range(len(images)):
                    number = str(i)
                    image_name = file_name.replace('.pdf', '.' + number + '.jpg')
                    images[i].save(image_name, 'JPEG')

def create_newspaper(exportdir, reportdir):
    """
    Build the newspaper based off of exported dashboards
    """

    reporttag  = REPORTTAG

    document = Document()

    header_section = document.sections[0]
    header = header_section.header
    header_text = header.paragraphs[0]
    header_text.text = 'Report: {} Generated: {}'.format(reporttag, NOWTIME)

    footer_section = document.sections[0]
    footer = footer_section.footer
    footer_text = footer.paragraphs[0]
    footer_text.text = 'Courtesy of SumoDashboardNews!'

    style = document.styles['Normal']
    font = style.font

    font.name = 'Calibri'
    font.size = Pt(14)

    counter = 1
    for key, value in DASHBOARDDICT.items():
        matchname = fnmatch.filter(os.listdir(exportdir), key + '.*.jpg')[0]
        if matchname is not None:
            jpgname = os.path.abspath(os.path.join(exportdir, matchname))
            picture_title = 'Title {}: {}'.format(counter, value)
            document.add_heading(picture_title, 0)
            document.add_picture(jpgname, width=Inches(4))
            document.add_paragraph('This is a sample summary for {}'.format(value))
            document.add_page_break()
            counter = counter + 1

    reportfile = '.'.join((reporttag, DATE, TIME, 'docx'))
    newspaper = os.path.abspath(os.path.join(reportdir, reportfile))

    if ARGS.verbose > 5:
        print('Printing Newspaper: {}'.format(newspaper))

    document.save(newspaper)

### class ###
class SumoApiClient():
    """
    This is defined SumoLogic API Client
    The class includes the HTTP methods, cmdlets, and init methods
    """
    def __init__(self, accessId=sumo_uid, accessKey=sumo_key, endpoint=None, \
                 caBundle=None, cookieFile='cookies.txt'):
        self.session = requests.Session()
        self.session.auth = (accessId, accessKey)
        self.default_version = 'v2'
        self.session.headers = {'content-type': 'application/json', 'accept': '*/*'}
        if caBundle is not None:
            self.session.verify = caBundle
        cookiejar = cookielib.FileCookieJar(cookieFile)
        self.session.cookies = cookiejar
        if endpoint is None:
            self.endpoint = self._get_endpoint()
        elif len(endpoint) < 3:
            self.endpoint = 'https://api.' + endpoint + '.sumologic.com/api'
        else:
            self.endpoint = endpoint
        if self.endpoint[-1:] == "/":
            raise Exception("Endpoint should not end with a slash character")

    def _get_endpoint(self):
        """
        SumoLogic REST API endpoint changes based on the geo location of the client.
        This method makes a request to the default REST endpoint and resolves the 401 to learn
        the right endpoint
        """
        self.endpoint = 'https://api.sumologic.com/api'
        self.response = self.session.get('https://api.sumologic.com/api/v1/collectors')
        endpoint = self.response.url.replace('/v1/collectors', '')
        return endpoint

    def get_versioned_endpoint(self, version):
        """
        formats and returns the endpoint and version
        """
        return self.endpoint+'/%s' % version

    def delete(self, method, params=None, version=None):
        """
        HTTP delete
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        response = self.session.delete(endpoint + method, params=params)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def get(self, method, params=None, version=None):
        """
        HTTP get
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        response = self.session.get(endpoint + method, params=params)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def get_file(self, method, params=None, version=None, headers=None):
        """
        HTTP get file
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        response = self.session.get(endpoint + method, params=params, headers=headers)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def post(self, method, params, headers=None, version=None):
        """
        HTTP post
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        response = self.session.post(endpoint + method, data=json.dumps(params), headers=headers)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def post_file(self, method, params, headers=None, version=None):
        """
        Handle file uploads via a separate post request to avoid clearing the content-type header
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        post_params = {'merge': params['merge']}
        file_data = open(params['full_file_path'], 'rb').read()
        files = {'file': (params['file_name'], file_data)}
        response = requests.post(endpoint + method, files=files, params=post_params,
                auth=(self.session.auth[0], self.session.auth[1]), headers=headers)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def put(self, method, params, headers=None, version=None):
        """
        HTTP put
        """
        version = version or self.default_version
        endpoint = self.get_versioned_endpoint(version)
        response = self.session.put(endpoint + method, data=json.dumps(params), headers=headers)
        if 400 <= response.status_code < 600:
            response.reason = response.text
        response.raise_for_status()
        return response

    def dashboards(self, monitors=False):
        """
        Return a list of dashboards
        """
        params = {'monitors': monitors}
        response = self.get('/dashboards', params)
        return json.loads(response.text)['dashboards']

    def dashboard(self, dashboard_id):
        """
        Return details on a specific dashboard
        """
        response = self.get('/dashboards/' + str(dashboard_id))
        return json.loads(response.text)['dashboard']

    def dashboard_data(self, dashboard_id):
        """
        Return data from a specific dashboard
        """
        response = self.get('/dashboards/' + str(dashboard_id) + '/data')
        return json.loads(response.text)['dashboardMonitorDatas']

    def export_dashboard(self,body):
        """
        Export data from a specific dashboard via a defined job
        """
        response = self.post('/dashboards/reportJobs', params=body, version='v2')
        job_id = json.loads(response.text)['id']
        if ARGS.verbose > 5:
            print('Started Job: {}'.format(job_id))
        return job_id

    def check_export_dashboard_status(self,job_id):
        """
        Check on the status a defined export job
        """
        response = self.get('/dashboards/reportJobs/%s/status' % (job_id), version='v2')
        response = {
            "result": json.loads(response.text),
            "job": job_id
        }
        return response

    def get_export_dashboard_result(self,job_id):
        """
        Retrieve the results of a defined export job
        """
        response = self.get_file(f"/dashboards/reportJobs/{job_id}/result", version='v2', \
                                 headers={'content-type': 'application/json', 'accept': '*/*'})
        response = {
            "job": job_id,
            "format": response.headers["Content-Type"],
            "bytes": response.content
        }
        if ARGS.verbose > 5:
            print ('Returned File Type: {}'.format(response['format']))
        return response

    def define_export_job(self,report_id,timezone="America/Los_Angeles",exportFormat='Pdf'):
        """
        Define a dashboard export job
        """
        payload = {
            "action": {
                "actionType": "DirectDownloadReportAction"
                },
            "exportFormat": exportFormat,
            "timezone": timezone,
            "template": {
                "templateType": "DashboardTemplate",
                "id": report_id
                }
        }
        return payload

    def poll_export_dashboard_job(self,job_id,tries=60,seconds=SLEEP):
        """
        Iterate and check on the dashboard export job
        """
        progress = ''
        tried=0

        while progress != 'Success' and tried < tries:
            tried += 1
            response = self.check_export_dashboard_status(job_id)
            progress = response['result']['status']
            time.sleep(seconds)

        if ARGS.verbose > 5:
            print('{}/{} job: {} status: {}'.format(tried, tries, \
                                                    job_id, response['result']['status']))
        response['tried'] = tried
        response['seconds'] = tried * seconds
        response['tries'] = tries
        response['max_seconds'] = tries * seconds
        return response

    def run_export_job(self,report_id,timezone="America/Los_Angeles", \
                       exportFormat='Pdf',tries=30,seconds=SLEEP):
        """
        Run the defined dashboard export job
        """
        payload = self.define_export_job(report_id,timezone=timezone,exportFormat=exportFormat)
        job = self.export_dashboard(payload)
        if ARGS.verbose > 7:
            print ('Running Job: {}'.format(job))
        poll_status = self.poll_export_dashboard_job(job,tries=tries,seconds=seconds)
        if poll_status['result']['status'] == 'Success':
            export = self.get_export_dashboard_result(job)
        else:
            print ('Job Unsuccessful after: {} attempts'.format(tries))
            export = {
                'job': job
            }
        export['id'] = report_id
        export['status'] = poll_status['result']['status']
        export['poll_status'] = poll_status
        return export

### class ###

if __name__ == '__main__':
    main()
